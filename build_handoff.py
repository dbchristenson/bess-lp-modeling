"""Generate a hand-off Word document for the team with methodology filled in
and remaining sections left blank."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_placeholder(doc, text="[To be completed]"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def add_body(doc, text):
    doc.add_paragraph(text)


def add_bold_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0, 0, 0)

    # ── Title ──
    title = doc.add_heading(
        "Battery Energy Storage for Data Center Grid Flexibility:\n"
        "A Quantitative Analysis of Dublin's I-SEM Market",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Table of Contents placeholder ──
    doc.add_heading("Table of Contents", level=1)
    add_placeholder(doc, "[Insert table of contents before final submission]")
    doc.add_page_break()

    # ── Executive Summary ──
    doc.add_heading("Executive Summary", level=1)
    add_placeholder(doc)
    doc.add_page_break()

    # ── Methodology ──
    doc.add_heading("Methodology", level=1)

    add_body(
        doc,
        "Our model evaluates whether a representative Google DeepMind data center "
        "in Dublin could reduce its electricity costs by installing a battery energy "
        "storage system (BESS) and participating in Ireland’s demand response "
        "programs. Using synthetic hourly electricity prices calibrated to real I-SEM "
        "market data, the model determines the optimal battery size and hour-by-hour "
        "charging schedule that minimizes the facility’s total annual electricity "
        "costs. The model and some inputs are taken and updated from a paper by Lukas "
        'Ljungblom, "Optimal Planning of Data Centers with On-Site Generation and '
        'Storage".',
    )

    add_bold_body(doc, "Three scenarios are compared:")

    table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    hdr.cells[0].text = "Scenario"
    hdr.cells[1].text = "Description"
    for cell in hdr.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    scenarios = [
        (
            "Grid-Only",
            "Baseline — data center draws a constant 10 MW from the grid with no battery",
        ),
        (
            "Grid + BESS",
            "Battery performs energy arbitrage: charges when electricity is cheap, "
            "discharges when it is expensive",
        ),
        (
            "Grid + BESS + DR",
            "Battery also earns revenue from EirGrid’s capacity market as a "
            "demand side unit (DSU)",
        ),
    ]
    for i, (name, desc) in enumerate(scenarios, start=1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = desc

    # Assumptions
    doc.add_heading("Assumptions", level=2)
    assumptions = [
        "The data center draws a constant 10 MW load and has a 20 MW grid connection, "
        "leaving headroom for battery charging.",
        "Google DeepMind workloads are not mission-critical, so the data center can "
        "curtail its load if the battery is unavailable.",
        "The BESS has a 15-year lifetime, 95% round-trip efficiency, and costs are "
        "discounted at an 8% weighted average cost of capital (WACC).",
        "The BESS capacity is always awarded in EirGrid’s capacity auction, making "
        "it eligible for annual demand side unit (DSU) payments up to 10 MW.",
    ]
    for a in assumptions:
        doc.add_paragraph(a, style="List Bullet")

    # Modeling
    doc.add_heading("Modeling", level=2)
    add_bold_body(
        doc,
        "The systems we modeled are easily divided into two categories: costs and revenues.",
    )

    # Costs
    doc.add_heading("Costs", level=3)
    add_body(
        doc,
        "The model accounts for four cost components that together form the facility’s "
        "annualized lifecycle cost (ALCC):",
    )
    costs = [
        "Grid subscription fee — a fixed annual charge of approximately €4.4M "
        "for maintaining the 10 MW grid connection, paid regardless of how much "
        "electricity is consumed.",
        "Energy import costs — variable costs for each MWh drawn from the grid, "
        "equal to the I-SEM spot price plus a time-of-use (TOU) network tariff. TOU "
        "rates are highest during weekday peak hours (17:00–19:00) and lowest "
        "overnight and on weekends.",
        "BESS capital expenditure (CapEx) — the upfront cost of the battery system, "
        "spread over its 15-year lifetime using a standard annuity factor at 8% WACC. "
        "CapEx scales with both the battery’s power rating (MW) and its energy "
        "capacity (MWh).",
        "BESS operating expenditure (OpEx) — annual maintenance and operating costs, "
        "also scaling with battery power and energy capacity. These costs are adjusted "
        "upward from thesis values to reflect inflation and rising labor costs in Ireland.",
    ]
    for i, c in enumerate(costs, start=1):
        doc.add_paragraph(c, style="List Number")

    add_body(
        doc,
        "For BESS cost assumptions, the model runs two cost scenarios: the original thesis "
        "estimates (NREL ATB 2025, Moderate) and the NREL ATB 2024 Advanced scenario, which "
        "reflects aggressive but plausible cost reductions in battery technology. All ATB costs "
        "are reported in 2022 USD and converted to EUR at a rate of 0.95 EUR/USD. The updated "
        "scenario is used for scenario comparison and policy recommendations.",
    )

    # Revenues and Savings
    doc.add_heading("Revenues and Savings", level=3)
    add_body(
        doc,
        "The model captures value through two mechanisms: energy arbitrage (cost reduction) "
        "and DSU capacity payments (revenue).",
    )
    add_body(
        doc,
        "Energy arbitrage is embedded in the BESS optimization itself. The battery charges "
        "when electricity prices are low — including during negative price events caused "
        "by wind curtailment — and discharges during expensive peak hours, reducing the "
        "facility’s net import costs. This benefit appears in all BESS scenarios automatically.",
    )
    add_body(
        doc,
        "Capacity payments are additional revenue available only in the DR scenario. EirGrid’s "
        "capacity auction market compensates large flexible loads that can reduce their grid draw "
        "during system stress events. These participants are called demand side units (DSUs). While "
        "DSUs can take several forms (curtailed loads, batteries, virtual power plants), our model "
        "registers only the battery — not the data center load itself — as the DSU. Revenue "
        "assumptions are based on the capacity-weighted average clearing price from EirGrid’s "
        "2025–2026 T-1 auction results.",
    )
    add_body(
        doc,
        "Additional detail on the mathematical formulation of the objective function and full "
        "model results can be found in the appendix and accompanying deliverables.",
    )

    # Emissions Reduction
    doc.add_heading("Emissions Reduction", level=3)
    add_body(
        doc,
        "Because the electricity grid delivers undifferentiated electrons, simply installing "
        "a battery does not change the carbon intensity of the power a data center consumes. "
        "However, during DR events the BESS allows the facility to withdraw from the grid "
        "entirely, removing demand that EirGrid would otherwise meet by dispatching a marginal "
        "peaking unit. In Ireland, the marginal generators called during system stress events "
        "are typically oil- or heavy-fuel-oil-fired (HFO) open-cycle plants — most notably "
        "Moneypoint, which operates as a generator of last resort.",
    )
    add_body(
        doc,
        "We quantify avoided emissions using a marginal displacement approach. For each DR "
        "event hour, the data center's full 10 MW load is served by the BESS rather than "
        "drawn from the grid. The emissions that would have been produced had that energy come "
        "from an HFO peaker (800 gCO₂/kWh) are compared against the emissions at Ireland's "
        "average grid carbon intensity (225 gCO₂/kWh, SEAI). The difference represents the "
        "system-level emissions avoided by displacing peaker generation.",
    )
    add_body(
        doc,
        "The peaker emission factor of 800 gCO₂/kWh is the midpoint of the 740–890 "
        "gCO₂/kWh range reported for oil-fired generation in the UNECE Life Cycle Assessment "
        "of Electricity Generation Options (2021). The grid-average figure is SEAI's published "
        "conversion factor for Irish electricity consumption.",
    )
    doc.add_page_break()

    # ── Results ──
    doc.add_heading("Results", level=1)
    add_placeholder(
        doc,
        "[To be completed — data tables, figures, and scenario comparisons "
        "will be provided by the modeler. See BESS_DR_Model_Results.xlsx and "
        "the figures/ directory for supporting material.]",
    )
    doc.add_page_break()

    # ── Regulatory and Policy Considerations ──
    doc.add_heading("Regulatory and Policy Considerations", level=1)
    add_placeholder(
        doc,
        "[To be completed — discuss regulatory implications for Ireland’s "
        "2030 Climate Action Plan, propose a policy or market mechanism "
        '(e.g., "Flexible Load Tariff" or "AI-as-Grid-Resource Pilot") '
        "that aligns corporate and national energy objectives.]",
    )
    doc.add_page_break()

    # ── Bibliography ──
    doc.add_heading("Bibliography", level=1)
    refs = [
        "Cole, W., & Karmakar, A. (2023). Cost projections for utility-scale battery "
        "storage: 2023 update (NREL/TP-6A40-85332). National Renewable Energy Laboratory. "
        "https://www.nrel.gov/docs/fy23osti/85332.pdf",
        "EirGrid. (n.d.). Demand side management. "
        "https://www.eirgrid.ie/industry/becoming-customer/demand-side-management",
        "EirGrid. (2024, September). EirGrid statement of charges 2024/2025 (v1.0). "
        "https://www.eirgrid.com",
        "Ljungblom, L. (2025). Optimal planning of data centers with on-site generation "
        "and storage [Master’s thesis, Chalmers University of Technology].",
        "National Renewable Energy Laboratory. (2025). Utility-scale battery storage. "
        "Annual Technology Baseline. "
        "https://atb.nrel.gov/electricity/2024/commercial_battery_storage",
        "SEAI. (n.d.). Conversion factors. SEAI Statistics. "
        "https://www.seai.ie/data-and-insights/seai-statistics/conversion-factors",
        "SEMO. (2025, July). Final capacity auction results: 2025/26 T-1 (FCAR2526T-1). "
        "https://www.sem-o.com/sites/semo/files/2025-07/FCAR2526T-1.pdf",
        "United Nations Economic Commission for Europe. (2021). Life cycle assessment of "
        "electricity generation options. UNECE. "
        "https://unece.org/sites/default/files/2021-11/LCA_final.pdf",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.add_page_break()

    # ── Appendix ──
    doc.add_heading("Appendix", level=1)
    add_placeholder(
        doc,
        "[To be completed — include objective function formulation, "
        "full model parameters, and supplementary figures.]",
    )

    out = "handoff_draft.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build_document()
